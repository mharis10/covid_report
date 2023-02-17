import pandas as pd
import numpy as np
from datetime import date
from docx import Document
import os
import glob
import zipfile

#read data file
data = pd.read_csv("COVID-19_Diagnostic_Laboratory_Testing__PCR_Testing__Time_Series.csv")

#variables to store maximum and minmumn +ve and -ve cases
maxp = 0
minp = np.inf
maxn = 0
minn = np.inf

def create_document(states):

  #set variables as global so they can be used outside the function	
  global maxp, minp, maxn, minn, smaxp, sminp, smaxn, sminn

  #convert date column to datetime type
  data['date']= pd.to_datetime(data['date'])
  #loop over each state
  for st in states:
    print(st)
    #get dataframe for particular, current state only
    state = data.loc[data['state_name'] == st]
    #get the last 5 recent dates for this state
    last_5_dates = state.date.unique()
    last_5_dates = last_5_dates[-5:]
    #new dataframe containing data for 5 recent most days only
    state = state.loc[state.apply(lambda x: x.date in last_5_dates, axis=1)]
    #sort in descending order by date
    state.sort_values(by=['date'], ascending = False)
    #final dataframe with specific columns mentioned
    state_final = state[['state_name', 'overall_outcome', 'date', 'new_results_reported', 'total_results_reported']]
    #pivot the dataframe in order to get separate columns for negative and positive cases
    state_table = pd.pivot_table(state, values = ['new_results_reported', 'total_results_reported'], index=['date','state_name'], columns='overall_outcome').reset_index()
    #convert multindex to single index and concatenate the index values
    state_table.columns = state_table.columns.get_level_values(0)+'_'+state_table.columns.get_level_values(1)

    #for states such as Marshall Islands, Inconclusive column does not exist so the if condition simply checks for that and assigns column names accordingly
    if 'new_results_reported_Inconclusive' in state_table.columns: 
      state_table = state_table.drop(['state_name_', 'new_results_reported_Inconclusive', 'total_results_reported_Inconclusive'], axis=1) 
    else:
      state_table = state_table.drop(['state_name_'], axis=1) 

    #convert column names to well-formatted with Upper Case and spaces  
    state_table.columns = ['Date', 'New Negative Tests Reported', 'New Positive Tests Reported',
                        'Total Negative Tests Reported', 'Total Positive Tests Reported']
    state_table = state_table[['Date','New Positive Tests Reported','New Negative Tests Reported',
                            'Total Positive Tests Reported','Total Negative Tests Reported']]
    #convert Date column to Date only, this removes the Timestamp from the Date column                        
    state_table['Date'] = state_table['Date'].dt.date
    #get statename
    statename = state['state_name'].iloc[0]
    #EXTRA FUNCTIONALITY
    #convert total +ve and -ve columns to list
    pos = state_table['Total Positive Tests Reported'].values.tolist()
    neg = state_table['Total Negative Tests Reported'].values.tolist()

    #get max and min from the list
    max_pos = max(pos)
    min_pos = min(pos)

    max_neg = max(neg)
    min_neg = min(neg)

    #Update variables accordingly in order to set max and min +ve and -ve values
    if max_pos> maxp:
      maxp = max_pos
      smaxp = state['state_name'].iloc[0]

    if min_pos< minp:
      minp = min_pos
      sminp = state['state_name'].iloc[0]    
    
    if max_neg> maxn:
      maxn = max_neg
      smaxn = state['state_name'].iloc[0]

    if min_neg< minn:
      minn = min_neg
      sminn = state['state_name'].iloc[0]
    #EXTRA FUNCTIONALITY END  

    #create document
    document = Document()

    #header
    string = statename + " " + str(date.today())
    section = document.sections[0]
    header = section.header
    header = header.paragraphs[0]
    header.text = string

    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = document.add_table(state_table.shape[0]+1, state_table.shape[1])

    # add the header rows.
    for j in range(state_table.shape[-1]):
        t.cell(0,j).text = state_table.columns[j]

    # add the rest of the data frame
    for i in range(state_table.shape[0]):
        for j in range(state_table.shape[-1]):
            t.cell(i+1,j).text = str(state_table.values[i,j])

    document.save(str(statename)+'.docx')

states = data['state_name'].unique()
create_document(states)
#current directory
directory = os.getcwd()
#print(directory)
path = directory + '/*.docx'
files = glob.glob(path)

#create zip file
ZipFile = zipfile.ZipFile("covid_state.zip", "w" )
#add word docs to zip file
for a in files:
    ZipFile.write(os.path.basename(a), compress_type=zipfile.ZIP_DEFLATED)
ZipFile.close()
#remove word docs, only the zip file remains
for f in files:
	os.remove(f)
#print extra functionality results
print("The highest number of positive cases are:", maxp, "in", smaxp)
print("The lowest number of positive cases are:", minp, "in", sminp)
print("The highest number of negative cases are:", maxn, "in", smaxn)
print("The lowest number of negative cases are:", minn, "in", sminn)